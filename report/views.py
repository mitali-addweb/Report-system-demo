from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.contrib.auth.decorators import login_required
from django.contrib.staticfiles import finders
from docx import Document
import weasyprint
import csv
from django.utils.encoding import smart_str
from django.contrib import messages
from django.core.files.storage import FileSystemStorage
from datetime import datetime

from .forms import ReportForm, AssetForm, ProblemForm
from .models import Report, Asset, ProblemType
from accounts.decorators import allowed_roles_or_permissions

# -------------------------
# Report Views
# -------------------------
@login_required
@allowed_roles_or_permissions(roles=['MDI Team'])
def create_report(request):
    if request.method == 'POST':
        form = ReportForm(request.POST)
        if form.is_valid():
            report = form.save(commit=False)
            report.author = request.user
            report.save()
            return redirect('report:create_report')
    else:
        form = ReportForm()

    return render(request, 'report/create_report.html', {'form': form})


@login_required
@allowed_roles_or_permissions(permissions=['report.view_report'])
def display_report(request):
    reports = Report.objects.all().order_by('-id')
    return render(request, 'report/display_report.html', {'reports': reports})


@login_required
@allowed_roles_or_permissions(roles=['MDI Team'], permissions=['report.view_report'])
def view_report(request, report_id):
    report = get_object_or_404(Report, id=report_id)
    return render(request, 'report/view_report.html', {'report': report})


@login_required
@allowed_roles_or_permissions(roles=['MDI Team'], permissions=['report.view_report'])
def update_report(request, report_id):
    report = get_object_or_404(Report, id=report_id)

    if request.method == 'POST':
        form = ReportForm(request.POST, instance=report)
        if form.is_valid():
            form.save()
            return redirect('report:view_report', report_id=report.id)
    else:
        form = ReportForm(instance=report)

    return render(request, 'report/update_report.html', {'form': form, 'report': report})


# -------------------------
# Export Report
# -------------------------
@login_required
@allowed_roles_or_permissions(roles=['MDI Team'], permissions=['report.view_report'])
def export_report_pdf(request, report_id):
    report = get_object_or_404(Report, id=report_id)
    html_string = render_to_string('report/report_pdf_template.html', {'report': report})

    css_file = finders.find('css/view.css')
    pdf_file = weasyprint.HTML(string=html_string).write_pdf(
        stylesheets=[weasyprint.CSS(css_file)]
    )

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Report_{report.id}.pdf"'
    return response

@login_required
@allowed_roles_or_permissions(roles=['MDI Team'], permissions=['report.view_report'])
def export_report_word(request, report_id):
    report = get_object_or_404(Report, id=report_id)
    doc = Document()

    def add_report_to_doc(rpt, level=1):
        """Recursive function to add report and its previous entries."""
        doc.add_heading(f'Report {rpt.id}', level=level)
        doc.add_paragraph(f'Asset ID: {rpt.asset_id}')
        doc.add_paragraph(f'Asset: {rpt.asset}')
        doc.add_paragraph(f'Asset Description: {rpt.asset.description}')
        doc.add_paragraph(f'Author: {rpt.author.username}')
        doc.add_paragraph(f'Entry Date: {rpt.entry_date}')
        doc.add_paragraph(f'Priority: {rpt.priority}')
        doc.add_paragraph(f'Status: {rpt.status}')
        doc.add_paragraph(f'Work Order Number: {rpt.work_order_number}')
        doc.add_paragraph(f'Problem Type: {rpt.problem_type.name}')
        doc.add_paragraph(f'Problem Description: {rpt.problem_description}')
        doc.add_paragraph(f'Recommended Action: {rpt.recommended_action}')
        doc.add_paragraph('')  # blank line for spacing

        # If previous entry exists, add a separate heading for it
        if rpt.previous_entry:
            doc.add_heading(f'Previous Report {rpt.previous_entry.id}', level=level)
            add_report_to_doc(rpt.previous_entry, level=level+1)

    add_report_to_doc(report)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=Report_{report.id}.docx'
    doc.save(response)
    return response

@login_required
@allowed_roles_or_permissions(roles=['MDI Team'], permissions=['report.view_report'])
def export_report_csv(request, report_id):
    report = get_object_or_404(Report, id=report_id)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename=Report_{report.id}.csv'

    writer = csv.writer(response)

    headers = [
        "Report ID", "Asset ID", "Asset Name", "Asset Description",
        "Author", "Entry Date", "Priority", "Status",
        "Work Order Number", "Problem Type",
        "Problem Description", "Recommended Action"
    ]
    writer.writerow(headers)

    report_chain = []
    current = report
    while current:
        report_chain.append(current)
        current = current.previous_entry

    report_chain.reverse()

    for rpt in report_chain:
        writer.writerow([
            rpt.id,
            rpt.asset.id,
            smart_str(rpt.asset.name),
            smart_str(rpt.asset.description or ""),
            rpt.author.username if rpt.author else "N/A",
            rpt.entry_date,
            rpt.priority,
            rpt.status,
            rpt.work_order_number or "",
            rpt.problem_type.name if rpt.problem_type else "",
            smart_str(rpt.problem_description),
            smart_str(rpt.recommended_action or "")
        ])

    return response


# -------------------------
# Import Report
# -------------------------

@login_required
@allowed_roles_or_permissions(roles=['MDI Team'])
def import_report_csv(request):
    if request.method == "POST" and "csv_file" in request.FILES:
        csv_file = request.FILES["csv_file"]

        if not csv_file.name.lower().endswith('.csv'):
            messages.error(request, "File is not CSV type")
            return redirect("report:import_report")

        file_data = csv_file.read().decode("utf-8").splitlines()
        reader = list(csv.reader(file_data))
        header = reader[0]
        rows = reader[1:]

        previous_report = None  

        for row in rows:
            try:
                asset_id, asset_name, asset_desc, author, entry_date, priority, status, work_order, problem_type, problem_desc, recommended_action = row

                asset, _ = Asset.objects.get_or_create(
                    id=int(asset_id),
                    defaults={"name": asset_name.strip(), "description": asset_desc.strip()}
                )

                problem_type_obj = None
                if problem_type.strip():
                    problem_type_obj, _ = ProblemType.objects.get_or_create(name=problem_type.strip())

                try:
                    entry_date_parsed = datetime.strptime(entry_date.strip(), "%m/%d/%Y").date()
                except ValueError:
                    entry_date_parsed = datetime.now().date()

                report = Report.objects.create(
                    asset=asset,
                    author=request.user,
                    entry_date=entry_date_parsed,
                    priority=priority.strip(),
                    status=status.strip(),
                    work_order_number=work_order.strip(),
                    problem_type=problem_type_obj,
                    problem_description=problem_desc.strip(),
                    recommended_action=recommended_action.strip(),
                    previous_entry=previous_report
                )

                previous_report = report

            except Exception as e:
                messages.error(request, f"Error on row {row}: {e}")
                continue

        messages.success(request, "CSV Imported Successfully")
        return redirect("report:display_report")

    return render(request, "report/import_report.html")


# -------------------------
# Asset Views
# -------------------------
@login_required
@allowed_roles_or_permissions(roles=['MDI Team'])
def create_asset(request):
    if request.method == 'POST':
        form = AssetForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('report:display_assets')
    else:
        form = AssetForm()

    return render(request, 'asset/create_asset.html', {'form': form})


@login_required
@allowed_roles_or_permissions(roles=['MDI Team'])
def display_assets(request):
    assets = Asset.objects.all().order_by('name')
    return render(request, 'asset/display_assets.html', {'assets': assets})


@login_required
@allowed_roles_or_permissions(roles=['MDI Team'])
def view_asset(request, asset_id):
    asset = get_object_or_404(Asset, id=asset_id)
    return render(request, 'asset/view_asset.html', {'asset': asset})


@login_required
@allowed_roles_or_permissions(roles=['MDI Team'])
def update_asset(request, asset_id):
    asset = get_object_or_404(Asset, id=asset_id)

    if request.method == 'POST':
        form = AssetForm(request.POST, instance=asset)
        if form.is_valid():
            form.save()
            return redirect('report:view_asset', asset_id=asset.id)
    else:
        form = AssetForm(instance=asset)

    return render(request, 'asset/update_asset.html', {'form': form, 'asset': asset})


# -------------------------
# Problem Type Views
# -------------------------
@login_required
@allowed_roles_or_permissions(roles=['MDI Team'])
def create_problem_type(request):
    if request.method == 'POST':
        form = ProblemForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('report:create_problem_type')
    else:
        form = ProblemForm()

    return render(request, 'problem_type/create_problem_type.html', {'form': form})


@login_required
@allowed_roles_or_permissions(roles=['MDI Team'])
def display_problem_type(request):
    problem_types = ProblemType.objects.all().order_by('name')
    return render(request, 'problem_type/display_problem_type.html', {'problem_types': problem_types})


@login_required
@allowed_roles_or_permissions(roles=['MDI Team'])
def update_problem_type(request, problem_type_id):
    problem_type = get_object_or_404(ProblemType, id=problem_type_id)

    if request.method == 'POST':
        form = ProblemForm(request.POST, instance=problem_type)
        if form.is_valid():
            form.save()
            return redirect('report:display_problem_type')
    else:
        form = ProblemForm(instance=problem_type)

    return render(request, 'problem_type/update_problem_type.html', {'form': form, 'problem_type': problem_type})
